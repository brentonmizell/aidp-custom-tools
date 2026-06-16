"""
Document Extract Tool
=====================
Extract clean text (and simple table text) from a document in a specific format:
PDF, DOCX, CSV, or plain text.

Three input modes (mutually exclusive — first non-empty wins):
  - `source_uri` — fetch the file bytes from an AIDP target (master volume or
    workspace) via the shared `aidp_io` module. Keeps extraction symmetric with
    the rest of the toolkits.
  - `content_base64` — raw base64-encoded bytes (works for PDF/DOCX/binary).
  - `content` — plain text/CSV pasted inline.

Optional sink:
  - `dest_uri` — when provided, the extracted text (or CSV rendering) is also
    written back to an AIDP target.

Format is taken from the `format` parameter, or guessed from the source URI / a
`filename`. The PDF/DOCX/CSV extractors all operate on BYTES wrapped in
io.BytesIO so pdfplumber, pypdf, and python-docx can read in-memory and don't
need a local file path.
"""

import base64
import io
import os

from aidputils.agents.tools.custom_tools.base import CustomToolBase
from .utils.config_utils import get_cfg

# Debug Channel — optional in runtime; provide no-op fallback so the tool still
# works if the runtime hasn't injected aidp_debug.
try:
    from aidp_debug import debug, debug_warn, debug_error, DebugLog  # type: ignore
except ImportError:  # pragma: no cover
    def debug(*args, **kwargs):
        pass

    def debug_warn(*args, **kwargs):
        pass

    def debug_error(*args, **kwargs):
        pass

    class DebugLog:  # type: ignore
        @staticmethod
        def embed(result):
            return result


# Shared file-IO module — optional. When absent (older runtimes, or local unit
# tests) source_uri / dest_uri simply produce a clear error envelope; the
# base64 / content paths still work.
try:
    from aidp_io import (  # type: ignore
        parse_uri as _aidp_parse_uri,
        read_file as _aidp_read_file,
        write_file as _aidp_write_file,
    )
    _AIDP_IO_AVAILABLE = True
except Exception:  # pragma: no cover
    _aidp_parse_uri = None
    _aidp_read_file = None
    _aidp_write_file = None
    _AIDP_IO_AVAILABLE = False


def _ok(data, **extra):
    """Build a success envelope while preserving legacy top-level keys."""
    payload = {"ok": True, "data": data}
    # Preserve legacy top-level keys for callers that read them directly.
    if isinstance(data, dict):
        for k, v in data.items():
            if k not in payload:
                payload[k] = v
    payload.update(extra)
    return payload


def _err(msg, error_type="ToolError", **extra):
    payload = {"ok": False, "error": msg, "error_type": error_type}
    payload.update(extra)
    return payload


def _guess_format_from_path(path):
    """Pull a file extension off a URI path or filename; return lowercase or ''."""
    if not path:
        return ""
    base = path.rsplit("/", 1)[-1]
    if "." not in base:
        return ""
    return base.rsplit(".", 1)[-1].lower()


@CustomToolBase.register
class DocumentExtractTool(CustomToolBase):
    """Extract text from a document (PDF, DOCX, CSV, or text).

    Provide ONE of:
      - `source_uri`    (master:<cat>.<sch>.<vol>:/path or workspace:/path)
      - `content_base64` (raw bytes, base64-encoded)
      - `content`       (inline text / CSV)

    Plus a `format` (pdf/docx/csv/txt/md/json) or a `filename` to infer it.
    Optionally writes the extracted text to `dest_uri` as well. Use to turn a
    document into text an agent can read or feed into RAG."""

    @classmethod
    def _execute_tool(cls, conf, runtime_params, **context_vars):
        fmt = (runtime_params.get("format", "") or "").lower().lstrip(".")
        filename = runtime_params.get("filename", "")
        source_uri = (runtime_params.get("source_uri") or "").strip()
        dest_uri = (runtime_params.get("dest_uri") or "").strip()

        # Infer format from source_uri or filename if not explicitly given.
        if not fmt and source_uri:
            fmt = _guess_format_from_path(source_uri).lstrip(".")
        if not fmt and filename and "." in filename:
            fmt = filename.rsplit(".", 1)[-1].lower()

        text_content = runtime_params.get("content", "")
        b64 = runtime_params.get("content_base64", "")

        max_chars = get_cfg(conf, "max_chars", 100000)
        max_csv_rows = get_cfg(conf, "max_csv_rows", 1000)
        max_input_bytes = get_cfg(conf, "max_input_bytes", 26214400)  # 25 MiB

        page_start = runtime_params.get("page_start")
        page_end = runtime_params.get("page_end")
        try:
            page_start = int(page_start) if page_start is not None else None
        except (TypeError, ValueError):
            page_start = None
        try:
            page_end = int(page_end) if page_end is not None else None
        except (TypeError, ValueError):
            page_end = None

        debug(
            "DocumentExtractTool start",
            format=fmt,
            filename=filename,
            source_uri=source_uri,
            dest_uri=dest_uri,
            has_b64=bool(b64),
            has_text=bool(text_content),
            page_start=page_start,
            page_end=page_end,
        )

        # ------------------------------------------------------------------ #
        # Resolve to bytes. Precedence: source_uri > content_base64 > content.
        # ------------------------------------------------------------------ #
        raw = None
        source = None
        if source_uri:
            if not _AIDP_IO_AVAILABLE:
                debug_error("aidp_io unavailable but source_uri provided")
                result = _err(
                    "source_uri requires the aidp_io module to be present in the runtime",
                    "DependencyError",
                )
                return DebugLog.embed(result)
            try:
                raw = _aidp_read_file(source_uri, conf, context_vars)
                if isinstance(raw, str):
                    raw = raw.encode("utf-8")
                source = "source_uri"
                debug("source_uri read", uri=source_uri, bytes=len(raw or b""))
            except Exception as e:
                debug_error("source_uri read failed", uri=source_uri, error=str(e))
                result = _err(
                    f"failed to read source_uri: {e}",
                    type(e).__name__,
                    source_uri=source_uri,
                )
                return DebugLog.embed(result)
        elif b64:
            try:
                raw = base64.b64decode(b64)
                source = "content_base64"
            except Exception as e:
                debug_error("base64 decode failed", error=str(e))
                result = _err("content_base64 is not valid base64", "InvalidInput")
                return DebugLog.embed(result)
        elif text_content:
            raw = text_content.encode("utf-8")
            source = "content"
        else:
            debug_warn("no content supplied")
            result = _err(
                "provide source_uri, content_base64 (binary), or content (text)",
                "InvalidInput",
            )
            return DebugLog.embed(result)

        # Byte cap on input.
        input_truncated = False
        if raw is not None and len(raw) > max_input_bytes:
            debug_warn(
                "input bytes exceeded cap; truncating",
                size=len(raw),
                cap=max_input_bytes,
            )
            raw = raw[:max_input_bytes]
            input_truncated = True

        if not fmt:
            debug_warn("format missing")
            result = _err(
                "format is required (pdf, docx, csv, or txt), or pass a filename / source_uri to infer it",
                "InvalidInput",
            )
            return DebugLog.embed(result)

        try:
            if fmt == "pdf":
                text, pages, page_count, pages_truncated = cls._extract_pdf(
                    raw, page_start, page_end
                )
                capped_text = text[:max_chars]
                data = {
                    "format": "pdf",
                    "pages": pages,
                    "page_count": page_count,
                    "chars": len(text),
                    "text": capped_text,
                    "truncated": len(text) > max_chars or pages_truncated or input_truncated,
                    "source": source,
                }
                cls._maybe_write_dest(dest_uri, capped_text, conf, context_vars, data)
                debug("pdf extracted", pages=pages, page_count=page_count, chars=len(text))
                return DebugLog.embed(_ok(data))

            if fmt in ("docx", "doc"):
                text, tables = cls._extract_docx(raw)
                capped_text = text[:max_chars]
                data = {
                    "format": "docx",
                    "tables": len(tables),
                    "chars": len(text),
                    "text": capped_text,
                    "table_text": tables[:50],
                    "truncated": len(text) > max_chars or len(tables) > 50 or input_truncated,
                    "source": source,
                }
                cls._maybe_write_dest(dest_uri, capped_text, conf, context_vars, data)
                debug("docx extracted", chars=len(text), tables=len(tables))
                return DebugLog.embed(_ok(data))

            if fmt == "csv":
                header, records, total_records = cls._extract_csv(raw)
                capped = records[:max_csv_rows]
                truncated = total_records > max_csv_rows or input_truncated
                data = {
                    "format": "csv",
                    "columns": header,
                    "row_count": total_records,
                    "rows": capped,
                    "returned_rows": len(capped),
                    "truncated": truncated,
                    "source": source,
                }
                # For CSV, dest_uri receives the raw decoded text (round-trippable).
                cls._maybe_write_dest(
                    dest_uri, raw.decode("utf-8", "replace"), conf, context_vars, data
                )
                debug(
                    "csv extracted",
                    columns=len(header),
                    row_count=total_records,
                    returned=len(capped),
                    truncated=truncated,
                )
                return DebugLog.embed(_ok(data))

            if fmt in ("txt", "text", "md", "json", "log"):
                text = raw.decode("utf-8", "replace")
                capped_text = text[:max_chars]
                data = {
                    "format": fmt,
                    "chars": len(text),
                    "text": capped_text,
                    "truncated": len(text) > max_chars or input_truncated,
                    "source": source,
                }
                cls._maybe_write_dest(dest_uri, capped_text, conf, context_vars, data)
                debug("text extracted", format=fmt, chars=len(text))
                return DebugLog.embed(_ok(data))

            debug_warn("unsupported format", format=fmt)
            result = _err(
                f"unsupported format '{fmt}'. Supported: pdf, docx, csv, txt/md/json.",
                "InvalidInput",
            )
            return DebugLog.embed(result)
        except Exception as e:
            debug_error("extraction failed", format=fmt, error=str(e))
            result = _err(str(e), type(e).__name__)
            return DebugLog.embed(result)

    # ------------------------------------------------------------------ #
    # Optional dest_uri writer — best-effort; surfaces errors via the data
    # envelope rather than failing the extraction.
    # ------------------------------------------------------------------ #
    @classmethod
    def _maybe_write_dest(cls, dest_uri, text_payload, conf, context_vars, data):
        if not dest_uri:
            return
        if not _AIDP_IO_AVAILABLE:
            data["dest_uri"] = dest_uri
            data["dest_write_error"] = (
                "aidp_io module not available in this runtime; dest_uri ignored"
            )
            debug_warn("dest_uri requested but aidp_io unavailable", uri=dest_uri)
            return
        try:
            payload_bytes = text_payload.encode("utf-8") if isinstance(text_payload, str) else bytes(text_payload)
            write_res = _aidp_write_file(dest_uri, payload_bytes, conf, context_vars)
            data["dest_uri"] = dest_uri
            data["dest_write"] = write_res
            debug("dest_uri write complete", uri=dest_uri, bytes=len(payload_bytes))
        except Exception as e:
            debug_error("dest_uri write failed", uri=dest_uri, error=str(e))
            data["dest_uri"] = dest_uri
            data["dest_write_error"] = str(e)

    # ------------------------------------------------------------------ #
    # Extractors — all accept raw BYTES wrapped in io.BytesIO so we never
    # need a temp file on disk.
    # ------------------------------------------------------------------ #
    @classmethod
    def _extract_pdf(cls, raw, page_start=None, page_end=None):
        """Extract text from a PDF, optionally restricted to a 1-based page range.

        Returns (text, pages_extracted, total_pages, range_truncated).
        Imports are lazy so callers using non-PDF formats don't pay the cost.
        Operates on io.BytesIO(raw) so source_uri bytes / base64 bytes / inline
        text bytes all flow through the same path.
        """
        # Prefer pdfplumber (better layout + tables); fall back to pypdf.
        try:
            import pdfplumber  # lazy
            with pdfplumber.open(io.BytesIO(raw)) as pdf:
                total = len(pdf.pages)
                start_idx, end_idx, truncated = cls._resolve_page_range(
                    total, page_start, page_end
                )
                parts = []
                for i in range(start_idx, end_idx):
                    parts.append(pdf.pages[i].extract_text() or "")
            return "\n\n".join(parts).strip(), (end_idx - start_idx), total, truncated
        except Exception as e:
            debug_warn("pdfplumber failed; falling back to pypdf", error=str(e))

        from pypdf import PdfReader  # lazy
        reader = PdfReader(io.BytesIO(raw))
        total = len(reader.pages)
        start_idx, end_idx, truncated = cls._resolve_page_range(
            total, page_start, page_end
        )
        parts = [(reader.pages[i].extract_text() or "") for i in range(start_idx, end_idx)]
        return "\n\n".join(parts).strip(), (end_idx - start_idx), total, truncated

    @staticmethod
    def _resolve_page_range(total, page_start, page_end):
        """Resolve 1-based inclusive page_start/page_end into 0-based [start, end).

        Returns (start_idx, end_idx, truncated_flag). truncated is True when the
        caller supplied a range that was clamped to fit the document.
        """
        if total <= 0:
            return 0, 0, False
        truncated = False
        start_idx = 0 if page_start is None else max(0, page_start - 1)
        end_idx = total if page_end is None else min(total, page_end)
        if page_start is not None and page_start - 1 > total - 1:
            truncated = True
            start_idx = total
        if page_end is not None and page_end > total:
            truncated = True
        if end_idx < start_idx:
            end_idx = start_idx
        return start_idx, end_idx, truncated

    @classmethod
    def _extract_docx(cls, raw):
        """Read a DOCX from BYTES via io.BytesIO — python-docx accepts any
        file-like object so we don't need a temp path."""
        from docx import Document  # lazy
        doc = Document(io.BytesIO(raw))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        tables = []
        for t in doc.tables:
            for row in t.rows:
                cells = [c.text.strip() for c in row.cells]
                if any(cells):
                    tables.append(" | ".join(cells))
        return "\n".join(paragraphs).strip(), tables

    @classmethod
    def _extract_csv(cls, raw):
        """Read a CSV from BYTES — decode then feed io.StringIO into csv.reader.

        Returns (header, records, total_record_count).
        """
        import csv

        rows = list(csv.reader(io.StringIO(raw.decode("utf-8", "replace"))))
        header = rows[0] if rows else []
        records = [dict(zip(header, r)) for r in rows[1:]] if header else []
        return header, records, len(records)
